import docx 

# Create a new Document
doc = docx.Document()

# Set the default font for the entire document to Times New Roman 12
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = docx.shared.Pt(12)

# Add a paragraph in Times New Roman bold
paragraph_bold = doc.add_paragraph()
run_bold = paragraph_bold.add_run("This is a bold paragraph.")
run_bold.font.bold = True
run_bold.font.name = 'Times New Roman'

# Add a paragraph in Times New Roman regular (not bold)
paragraph_regular = doc.add_paragraph()
run_regular = paragraph_regular.add_run("This is a regular paragraph.")
run_regular.font.bold = False  # Not specifying this line also works, as the default is False
run_regular.font.name = 'Times New Roman'

# Save the document
doc.save('formatted_document.docx')
