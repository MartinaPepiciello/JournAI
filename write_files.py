import io

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import Paragraph

import docx


def write_pdf(title, entries, prompts):
    # Create a BytesIO object to store the PDF
    output = io.BytesIO()

    # Create a PDF document
    c = canvas.Canvas(output, pagesize=letter)

    # Define the width and height for the text box
    width, height = letter

    # Set up the starting position for the text
    x, y = 50, height - 50

    # Create paragraph styles
    justified_normal = ParagraphStyle(
        name='JustifyNormal',
        fontName='Times-Roman',
        fontSize=12,
        alignment=4  # 4 represents 'Justify'
    )
    justified_bold = ParagraphStyle(
        name='JustifyBold',
        fontName='Times-Bold',
        fontSize=12,
        alignment=4
    )

    title_style = ParagraphStyle(
        name='Title',
        fontName='Times-Bold',
        fontSize=16,
        alignment=4
    )

    # Write title
    if title:
        x, y = add_pdf_paragraphs(title, title_style, c, x, y, width, height)
        y -= 30

    # Write prompts and entries in the PDF
    for prompt, entry in zip(prompts, entries):
        if prompt:
            x, y = add_pdf_paragraphs(prompt, justified_bold, c, x, y, width, height)
        x, y = add_pdf_paragraphs(entry, justified_normal, c, x, y, width, height)
        y -= 20


    # Save the PDF content
    c.save()
    output.seek(0)

    return output


def write_docx(title, entries, prompts):
    # Create a BytesIO object to store the document in memory
    output = io.BytesIO()

    # Create Document object
    doc = docx.Document()

    # Add title
    if title:
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.font.name = 'Times New Roman'
        title_run.font.size = docx.shared.Pt(16)
        title_run.bold = True
        doc.add_paragraph()

    # Regular font style
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = docx.shared.Pt(12)

    # Write prompts and entries as paragraphs in the document
    for prompt, entry in zip(prompts, entries):
        if prompt:
            add_docx_paragraphs(prompt, doc, True)
        add_docx_paragraphs(entry, doc, False)
        doc.add_paragraph()

    # Save the document to the in-memory buffer
    doc.save(output)
    output.seek(0)

    return output


def write_txt(title, entries, prompts):
    # Create the content for the text file
    content = ''
    if title:
        content += title.upper() + '\n\n'
    for prompt, entry in zip(prompts, entries):
        content += (prompt + '\n') if prompt else ''
        content += entry + '\n\n'

    # Create a BytesIO object to store the text file
    output = io.BytesIO()
    output.write(content.encode('utf-8'))
    output.seek(0)

    return output


# Helper function to write paragraphs in pdf's
def add_pdf_paragraphs(text, style, c, x_start, y_start, width, height):
    x = x_start
    y = y_start

    paragraphs = text.split('\n')

    for paragraph in paragraphs:
        if not paragraph.strip():
            continue
        
        # Create a Paragraph object with the text and the justified style
        p = Paragraph(paragraph, style)

        # Calculate the space left on the page
        space_left = y - 50  # margin at the bottom of the page
        text_height = p.wrap(width - 100, space_left)[1]

        # Check if there's enough space for the current paragraph
        if text_height > space_left:
            c.showPage()
            y = height - 50  # Reset y position

        p.drawOn(c, x, y - text_height)  # Draw the paragraph on the canvas
        y -= text_height + 10  # Adjust vertical position for the next paragraph

    return x, y

# Helper function to write paragraphs in docx's
def add_docx_paragraphs(text, doc, bold):
    paragraphs = text.split('\n')

    for paragraph in paragraphs:
        if paragraph.strip():
            p = doc.add_paragraph()
            my_run = p.add_run(paragraph.strip())
            my_run.font.bold = bold
